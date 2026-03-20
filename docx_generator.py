import sys, os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.append(os.path.dirname(SCRIPT_DIR))

import json
import pandas as pd
from docx import Document # type: ignore
from docx.shared import Pt, Cm # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # type: ignore
from utils.llm_call_functions import get_llm_output
from langchain_openai import ChatOpenAI # type: ignore
from utils.prompts import prompt_business_info, prompt_business_attribute_description


class ReportGenerator:
    """
    Generates a detailed report for a service or product analysis.
    """

    def __init__(self, result_dir):
        """
        Initializes the ReportGenerator instance.

        Args:
            result_dir (str): Directory to store result files.
        """
        self.model = ChatOpenAI(temperature=0, model="gpt-4o")
        self.result_dir = result_dir
        self.docx_file = 'outputs/docx/output.docx'
        self.doc = Document()
        self.img_index = 1
        self.set_default_style()

    def set_default_style(self):
        """
        Sets the default font style for the document.
        """
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times-Roman'
        font.size = Pt(12)

    def add_paragraph(self, text, style=None, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        """
        Adds a paragraph with specified text and style to the document.

        Args:
            text (str): Text content of the paragraph.
            style (dict, optional): Style dictionary for the paragraph.
            alignment (int, optional): Alignment of the paragraph.
        """
        paragraph = self.doc.add_paragraph(text)
        if style:
            run = paragraph.runs[0]
            run.font.size = style.get('size', Pt(14))
            run.bold = style.get('bold', False)
        paragraph.alignment = alignment

    def add_title(self, title, style):
        """
        Adds a title with specified style to the document.

        Args:
            title (str): Title text.
            style (dict): Style dictionary for the title.
        """
        self.add_paragraph(title, style, WD_PARAGRAPH_ALIGNMENT.CENTER)

    def add_section_header(self, header):
        """
        Adds a section header to the document.

        Args:
            header (str): Header text.
        """
        header_style = {'size': Pt(18), 'bold': True}
        self.add_paragraph(header, header_style)

    def create_business_description(self, business_info_dict, aquaventure_url):
        """
        Creates the business description for the report.

        Args:
            business_info_dict (dict): Information about the business.
            aquaventure_url (str): URL for additional business information.

        Returns:
            str: Generated business description.
        """
        business_info_text = ""
        for key in business_info_dict:
            business_info_text += f"{key}: {business_info_dict[key]}\n"
        description = get_llm_output(
            model=self.model,
            prompt=prompt_business_info,
            input_dict={"info": business_info_text},
            output_type="string"
        )
        description += f'\nThe data has been gathered from this [link]({aquaventure_url})'
        return description

    def add_user_scores(self, df):
        """
        Adds user scores section to the document.

        Args:
            df (pd.DataFrame): DataFrame containing user scores.
        """
        self.add_paragraph('Here we provide charts and statistics to better understand the scores given by the users')

        stats = [
            ["Statistic", "Value"],
            ["Mean", round(df['Score'].mean(), 2)],
            ["Median", round(df['Score'].median(), 2)],
            ["Standard Deviation", round(df['Score'].std(), 2)],
            ["Min", df['Score'].min()],
            ["Max", df['Score'].max()]
        ]

        table = self.doc.add_table(rows=len(stats), cols=len(stats[0]))
        hdr_cells = table.rows[0].cells

        for i in range(len(stats[0])):
            hdr_cells[i].text = stats[0][i]
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for row_idx in range(1, len(stats)):
            row_cells = table.rows[row_idx].cells
            for col_idx in range(len(stats[row_idx])):
                row_cells[col_idx].text = str(stats[row_idx][col_idx])
                row_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table.style = 'Table Grid'

        self.doc.add_paragraph()
        self.doc.add_picture(os.path.join(self.result_dir, 'userScore_UserScores_linechart.png'), width=Cm(15))
        self.add_paragraph(f"{self.img_index}. The line chart of scores given by users over time", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.img_index += 1

        self.doc.add_paragraph()
        self.doc.add_picture(os.path.join(self.result_dir, 'UserScore_Histogram.png'), width=Cm(15))
        self.add_paragraph(f"{self.img_index}. The histogram of scores given by users over time", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.img_index += 1

    def add_attributes_section(self, business_info_text, attributes):
        """
        Adds attributes section to the document.

        Args:
            business_info_text (str): Business information text.
            attributes (list): List of attributes related to the business.
        """
        attributes_description = (
            "In our assessment of user reviews, we have selected several characteristics "
            "pertinent to waterpark analysis. This section elaborates on these characteristics "
            "and their connection to our evaluation."
        )
        attributes_description += "\n" + get_llm_output(
            model=self.model,
            prompt=prompt_business_attribute_description,
            input_dict={"info": business_info_text, "attributes": attributes},
            output_type="string"
        )
        attributes_description += "\n\nIn the next section, we will analyze the mentioned attributes specifically."

        for line in attributes_description.split("\n"):
            self.add_paragraph(line)

        self.doc.add_paragraph()
        self.doc.add_picture(os.path.join(self.result_dir, 'allAttributesAverageScore.png'), width=Cm(15))
        self.add_paragraph(f"{self.img_index}. bar chart for the average score given to each attribute", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.img_index += 1

        self.doc.add_page_break()

    def add_attribute_summaries(self, attributes, summary_df, file_list):
        """
        Adds attribute summaries section to the document.

        Args:
            attributes (list): List of attributes related to the business.
            summary_df (pd.DataFrame): DataFrame containing attribute summaries.
            file_list (list): List of image files related to attributes.
        """
        for attribute in attributes:
            try:
                self.add_paragraph(attribute, {'size': Pt(18), 'bold': True})
                attr_row = summary_df.loc[summary_df['Attribute'] == attribute]
                attribute_summary = attr_row['Summary'].values[0]
                for line in attribute_summary.split("\n"):
                    self.add_paragraph(line)
                imgs_name = [name for name in file_list if attribute in name]
                # Line chart
                img_name = [name for name in imgs_name if "linechar" in name][0]
                img_path = os.path.join(self.result_dir, img_name)
                self.doc.add_picture(img_path, width=Cm(15))
                self.add_paragraph(f"{self.img_index}. Chart for {attribute} over time", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
                self.img_index += 1
                self.doc.add_paragraph()
                # Histogram
                img_name = [name for name in imgs_name if "hist.png" in name][0]
                img_path = os.path.join(self.result_dir, img_name)
                self.doc.add_picture(img_path, width=Cm(15))
                self.add_paragraph(f"{self.img_index}. Histogram for {attribute} over time", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
                self.img_index += 1

                self.add_paragraph('_' * 72)
            except Exception as e:
                print(f"Error adding {attribute}: {e}")

    def add_keyword_analysis(self):
        """
        Adds keyword analysis section to the document.
        """
        self.add_section_header('Keyword Analysis')
        description = "Charts below are created to show the amount and the number of words used in positive and negative reviews."
        self.add_paragraph(description)

        self.doc.add_picture(os.path.join(self.result_dir, 'PositiveWordCloud.png'), width=Cm(15))
        self.add_paragraph(f"{self.img_index}. Positive Keywords", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.img_index += 1

        self.doc.add_paragraph()
        self.doc.add_picture(os.path.join(self.result_dir, 'negativeWordCloud.png'), width=Cm(15))
        self.add_paragraph(f"{self.img_index}. Negative Keywords", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.img_index += 1

        self.doc.add_page_break()

    def add_semantic_analysis(self):
        """
        Adds semantic analysis section to the document.
        """
        self.add_section_header('Semantic Analysis')
        description = "The chart below shows the percentage of negative, positive and mixed reviews"
        self.add_paragraph(description)

        self.doc.add_picture(os.path.join(self.result_dir, 'SentimentPieChart.png'), width=Cm(15))
        self.add_paragraph(f"{self.img_index}. semantic pie chart", {'size': Pt(12), 'bold': False}, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.img_index += 1

        self.doc.add_page_break()

    def add_cost_report(self):
        """
        Adds cost report section to the document.
        """
        self.add_section_header('Cost Report')
        description = "At the end we have provided the cost of creating this report"
        self.add_paragraph(description)

        with open(os.path.join(self.result_dir, "PriceReport.txt")) as f:
            price_report = f.read().splitlines()
        for line in price_report:
            self.add_paragraph(line)

    def save_document(self):
        """
        Saves the generated document.
        """
        self.doc.save(self.docx_file)
        if os.path.exists(self.docx_file):
            print(f"The file '{self.docx_file}' has been successfully created.")
        else:
            print(f"Failed to create the file '{self.docx_file}'.")

    def create_report(self, business_info_dict, aquaventure_url, df, attributes, summary_df, file_list):
        """
        Creates the full report document.

        Args:
            business_info_dict (dict): Information about the business.
            aquaventure_url (str): URL for additional business information.
            df (pd.DataFrame): DataFrame containing user scores.
            attributes (list): List of attributes related to the business.
            summary_df (pd.DataFrame): DataFrame containing attribute summaries.
            file_list (list): List of image files related to attributes.
        """
        self.add_title('Aquaventure Waterpark Analysis Report', {
            'size': Pt(24),
            'bold': True
        })
        description = self.create_business_description(business_info_dict, aquaventure_url)

        self.add_section_header('Introduction')
        for line in description.split("\n"):
            self.add_paragraph(line)
        self.add_paragraph('_' * 72)
        self.add_paragraph('')

        self.add_section_header('Scores Given By Users')
        self.add_user_scores(df)

        self.add_section_header('Attributes of Reviews')
        self.add_attributes_section(description, attributes)
        self.add_attribute_summaries(attributes, summary_df, file_list)

        self.add_keyword_analysis()
        self.add_semantic_analysis()
        self.add_cost_report()

        self.save_document()


if __name__ == "__main__":
    result_dir = "AquaventureResult"  # This can be changed to any other directory as needed
    aquaventure_url = 'https://www.google.com.pk/travel/entity/key/ChoI3eTz_5vPx8e3ARoNL2cvMTFmOHAzMXF5NRAE/reviews?utm_campaign=sharing&utm_medium=link&utm_source=htls&ts=CAEaBAoCGgAqBAoAGgA'
    df = pd.read_csv(os.path.join(result_dir, "resultCSV.csv"))
    with open("config/product_review_criteria.json") as f:
        data = json.load(f)
    attributes = data["Waterpark"]
    summary_df = pd.read_csv(os.path.join(result_dir, "AttributesSummary.csv"))
    file_list = [file_name for file_name in os.listdir(result_dir) if file_name.endswith(('.jpg', '.png')) and "attribute" in file_name.lower()]

    # Define general_review and business_info_dict with placeholder values
    general_review = "World's largest water park with over 105 slides, attractions & experiences, plus a splash area for kids."
    business_info_dict = {"name": "Aquaventure", "type": "Waterpark", "about": general_review}

    report = ReportGenerator(result_dir)
    report.create_report(business_info_dict, aquaventure_url, df, attributes, summary_df, file_list)
